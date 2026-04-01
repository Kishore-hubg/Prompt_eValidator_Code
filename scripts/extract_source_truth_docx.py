from __future__ import annotations

import json
import re
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List

from docx import Document


@dataclass
class ParsedDoc:
    file_name: str
    file_path: str
    section: str
    paragraphs: List[str]


LEVEL_WEIGHTS = {
    "CRITICAL": 18,
    "HIGH": 14,
    "MEDIUM-HIGH": 10,
    "MEDIUM": 7,
    "LOW": 4,
}

PERSONA_SECTION_ANCHORS = {
    "persona_0": "Any Infovision employee using an AI-powered tool, chat interface, IDE copilot, or internal AI product",
    "persona_1": "Largest headcount cluster in Infovision; primary target for Phase 1 validator rollout",
    "persona_2": "Manages scrum-based delivery cycles across sprint planning, review, and retrospective events",
    "persona_3": "Business Intelligence cluster covering two business-facing personas with overlapping but distinct prompt needs",
    "persona_4": "Handles inbound and outbound customer and internal support tickets at high volume",
}


def _read_docx(path: Path) -> List[str]:
    doc = Document(str(path))
    lines: List[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)
    return lines


def parse_source_of_truth(base_dir: Path) -> List[ParsedDoc]:
    source_dir = base_dir / "docs" / "source_of_truth"
    parsed: List[ParsedDoc] = []
    for file_path in sorted(source_dir.rglob("*.docx")):
        section = file_path.parent.name
        paragraphs = _read_docx(file_path)
        parsed.append(
            ParsedDoc(
                file_name=file_path.name,
                file_path=str(file_path.relative_to(base_dir)).replace("\\", "/"),
                section=section,
                paragraphs=paragraphs,
            )
        )
    return parsed


def _normalize_dimension_name(label: str) -> str:
    cleaned = re.sub(r"[^a-z0-9]+", "_", label.lower()).strip("_")
    alias = {
        "context_completeness": "context",
        "ambiguity_reduction": "ambiguity_reduction",
        "constraint_definition": "constraints",
        "accuracy_orientation": "accuracy",
        "technical_precision": "technical_precision",
        "edge_case_coverage": "edge_cases",
        "output_format": "output_format",
        "business_relevance": "business_relevance",
        "retrieval_grounding": "grounding",
        "tone": "tone_empathy",
        "compliance_policy": "compliance",
        "structure": "output_format",
    }
    return alias.get(cleaned, cleaned)


def _collect_between(lines: List[str], start_marker: str, stop_markers: List[str]) -> List[str]:
    start_idx = -1
    for idx, line in enumerate(lines):
        if start_marker in line:
            start_idx = idx + 1
            break
    if start_idx == -1:
        return []

    out: List[str] = []
    for line in lines[start_idx:]:
        if any(marker in line for marker in stop_markers):
            break
        if line.strip():
            out.append(line.strip())
    return out


def _extract_dimension_weights(section_lines: List[str]) -> Dict[str, int]:
    rows = _collect_between(
        section_lines,
        "Quality Dimension Importance",
        ["What the Validator Must Check", "Clear task framing", "Most Useful Real-Time Suggestions"],
    )
    weights: Dict[str, int] = {}
    for row in rows:
        match = re.match(r"^([A-Za-z0-9 /&\-\(\)\.]+):\s*(CRITICAL|HIGH|MEDIUM-HIGH|MEDIUM|LOW)\b", row)
        if not match:
            continue
        dim_name = _normalize_dimension_name(match.group(1))
        weights[dim_name] = LEVEL_WEIGHTS[match.group(2)]

    return weights


def _extract_suggestions(section_lines: List[str]) -> List[str]:
    rows = _collect_between(
        section_lines,
        "Most Useful Real-Time Suggestions",
        ["Cluster Overview", "Role Responsibilities", "Appendix", "Quality Dimension Importance"],
    )
    suggestions: List[str] = []
    for row in rows:
        text = row.strip().strip('"')
        if text:
            suggestions.append(text)
    return suggestions


def _extract_checks(section_lines: List[str]) -> List[str]:
    checks = _collect_between(
        section_lines,
        "What the Validator Must Check",
        ["Penalty Triggers", "Most Useful Real-Time Suggestions", "Cluster Overview", "Role Responsibilities"],
    )
    if checks:
        return checks
    return _collect_between(
        section_lines,
        "Shared Cluster Checks",
        ["Penalty Triggers", "Most Useful Real-Time Suggestions", "Cluster Overview", "Role Responsibilities"],
    )


def _extract_penalties(section_lines: List[str]) -> List[str]:
    return _collect_between(
        section_lines,
        "Penalty Triggers",
        ["Most Useful Real-Time Suggestions", "Cluster Overview", "Role Responsibilities"],
    )


def _build_keyword_checks(checks: List[str], penalties: List[str], persona_id: str) -> Dict[str, List[str]]:
    base_terms = {
        "persona_0": {
            "output_format": ["bullet", "table", "paragraph", "template", "checklist", "numbered"],
            "constraints": ["tone", "word count", "scope", "audience", "include", "exclude"],
            "accuracy": ["source", "based on", "citation", "grounded", "do not infer"],
            "actionability": ["next step", "owner", "action item", "recommend"],
        },
        "persona_1": {
            "technical_precision": ["language", "framework", "version", "api", "codebase"],
            "edge_cases": ["edge case", "boundary", "null", "timeout", "error state", "concurrency"],
            "testability": ["test", "assert", "acceptance criteria", "expected result", "verify"],
            "role_alignment": ["new code", "bug fix", "refactor", "qa", "sdet", "automation"],
        },
        "persona_2": {
            "output_format": ["status report", "risk log", "sprint summary", "estimate comparison", "escalation note"],
            "context": ["sprint", "team", "velocity", "dates", "scope"],
            "prioritization": ["priority", "rank", "likelihood", "impact", "top 3", "moscow", "wsjf"],
            "traceability": ["history", "comparable story", "reference data", "baseline"],
        },
        "persona_3": {
            "grounding": ["source document", "section", "citation", "do not infer", "explicitly stated"],
            "traceability": ["source section", "story id", "reference", "trace"],
            "output_format": ["brd", "user story", "roadmap", "matrix", "process map"],
            "prioritization": ["moscow", "wsjf", "value vs effort", "priority"],
        },
        "persona_4": {
            "tone_empathy": ["empathetic", "firm", "apologetic", "informative", "neutral", "professional"],
            "compliance": ["policy", "sla", "refund", "regulatory", "compliant"],
            "output_format": ["email", "live chat", "phone script", "internal note", "closure", "escalation"],
            "speed": ["concise", "short", "quick", "fast"],
        },
    }

    merged_text = " ".join(checks + penalties).lower()
    result = base_terms.get(persona_id, {}).copy()
    for dim, terms in list(result.items()):
        hit_terms = [term for term in terms if term in merged_text]
        if hit_terms:
            result[dim] = hit_terms
    return result


def _find_section_start(lines: List[str], marker: str) -> int:
    for idx, line in enumerate(lines):
        if marker in line:
            return idx
    return -1


def _slice_persona_sections(v11_lines: List[str]) -> Dict[str, List[str]]:
    starts: List[tuple[str, int]] = []
    for persona_id, marker in PERSONA_SECTION_ANCHORS.items():
        idx = _find_section_start(v11_lines, marker)
        if idx != -1:
            starts.append((persona_id, idx))

    starts.sort(key=lambda x: x[1])
    sections: Dict[str, List[str]] = {}
    for i, (persona_id, start) in enumerate(starts):
        end = starts[i + 1][1] if i + 1 < len(starts) else len(v11_lines)
        sections[persona_id] = v11_lines[start:end]
    return sections


def generate_persona_criteria_from_v11(parsed_docs: List[ParsedDoc], base_dir: Path) -> Dict[str, dict] | None:
    v11_doc = None
    for item in parsed_docs:
        if "Infovision_Prompt_Validator_v11" in item.file_name:
            v11_doc = item
            break
    if not v11_doc:
        return None

    existing_path = base_dir / "app" / "config" / "persona_criteria_source_truth.json"
    if existing_path.exists():
        criteria = json.loads(existing_path.read_text(encoding="utf-8"))
    else:
        criteria = {}

    sections = _slice_persona_sections(v11_doc.paragraphs)

    for persona_id in ["persona_0", "persona_1", "persona_2", "persona_3", "persona_4"]:
        section_lines = sections.get(persona_id, [])
        if not section_lines:
            continue
        weights = _extract_dimension_weights(section_lines)
        checks = _extract_checks(section_lines)
        penalties = _extract_penalties(section_lines)
        suggestions = _extract_suggestions(section_lines)

        entry = criteria.get(persona_id, {"id": persona_id})
        if weights:
            entry["weights"] = weights
        if suggestions:
            entry["suggestions"] = suggestions[:8]
        entry["validator_checks"] = checks
        entry["penalty_triggers"] = penalties
        entry["keyword_checks"] = _build_keyword_checks(checks, penalties, persona_id)
        criteria[persona_id] = entry

    return criteria


def generate_prompt_guidelines_config(parsed_docs: List[ParsedDoc]) -> dict:
    prompt_docs = [d for d in parsed_docs if d.section == "prompt_guidelines"]
    return {
        "strict_mode": True,
        "strict_penalty_per_miss": 3,
        "strict_penalty_cap": 15,
        "claude_friendly_mode": True,
        "sources": [
            {"file_name": d.file_name, "file_path": d.file_path}
            for d in prompt_docs
        ],
        "frontier_model_applicability": {
            "anthropic_claude": [
                "clear_direct_action",
                "context_grounding",
                "output_format_control",
                "constraints_present",
                "structured_for_complexity",
                "factual_grounding_for_analysis",
            ],
            "google_gemini": [
                "clear_direct_action",
                "context_grounding",
                "output_format_control",
                "constraints_present",
                "structured_for_complexity",
                "factual_grounding_for_analysis",
            ],
            "perplexity": [
                "clear_direct_action",
                "context_grounding",
                "output_format_control",
                "constraints_present",
                "factual_grounding_for_analysis",
            ],
        },
        "claude_curated_priorities": [
            "Be clear and direct with explicit actions.",
            "Always include context and intended audience.",
            "Specify output format and constraints.",
            "For analytical prompts, require source-grounded output and citations.",
            "Use structure for complex prompts (steps, sections, tags).",
        ],
        "global_checks": [
            {
                "id": "clear_direct_action",
                "description": "Prompt starts with a clear action verb and direct instruction.",
                "keywords": ["summarize", "draft", "extract", "analyze", "explain", "create", "generate", "rewrite", "review"],
                "issue_if_missing": "Start with a clear action verb and direct task instruction."
            },
            {
                "id": "context_grounding",
                "description": "Prompt includes context, source, or motivation for the task.",
                "keywords": ["based on", "context", "source", "document", "given", "from", "audience", "purpose"],
                "issue_if_missing": "Add context about source material, audience, and purpose."
            },
            {
                "id": "output_format_control",
                "description": "Prompt declares expected output format and structure.",
                "keywords": ["bullet", "table", "paragraph", "json", "yaml", "xml", "numbered", "steps", "template"],
                "issue_if_missing": "Specify the required output format and structure."
            },
            {
                "id": "constraints_present",
                "description": "Prompt defines constraints such as tone, length, scope, or exclusions.",
                "keywords": ["tone", "word", "limit", "scope", "include", "exclude", "must", "should", "do not"],
                "issue_if_missing": "Add explicit constraints (tone, length, scope, inclusions/exclusions)."
            },
            {
                "id": "structured_for_complexity",
                "description": "Complex prompts should use structure (steps, tags, or sections).",
                "keywords": ["step", "1.", "2.", "<", "section", "format", "schema"],
                "issue_if_missing": "For complex prompts, use explicit structure (steps/sections/tags).",
                "min_word_count": 25
            },
            {
                "id": "factual_grounding_for_analysis",
                "description": "Analytical extraction tasks should request grounded/cited output.",
                "applies_when_any": ["analyze", "extract", "compare", "summarize", "requirements"],
                "keywords": ["based on", "source", "cite", "citation", "only explicit", "do not infer"],
                "issue_if_missing": "For analysis/extraction tasks, require grounded output with source references."
            }
        ]
    }


def main() -> None:
    base_dir = Path(__file__).resolve().parent.parent
    output_dir = base_dir / "app" / "config"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "source_of_truth_extracted.json"

    parsed = parse_source_of_truth(base_dir)

    payload = {
        "source_dir": "docs/source_of_truth",
        "documents_found": len(parsed),
        "documents": [asdict(item) for item in parsed],
    }
    output_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")

    persona_criteria = generate_persona_criteria_from_v11(parsed, base_dir)
    persona_path = output_dir / "persona_criteria_source_truth.json"
    guidelines_path = output_dir / "prompt_guidelines_source_truth.json"
    guidelines = generate_prompt_guidelines_config(parsed)

    if persona_criteria:
        persona_path.write_text(json.dumps(persona_criteria, indent=2), encoding="utf-8")
    guidelines_path.write_text(json.dumps(guidelines, indent=2), encoding="utf-8")

    if parsed:
        print(f"Extracted {len(parsed)} document(s) -> {output_path}")
        if persona_criteria:
            print(f"Generated persona criteria -> {persona_path}")
        else:
            print("Skipped persona criteria generation (v11 document not found).")
        print(f"Generated prompt guidelines config -> {guidelines_path}")
    else:
        print(
            "No DOCX files found under docs/source_of_truth. "
            "Please verify files are inside this workspace path."
        )


if __name__ == "__main__":
    main()
